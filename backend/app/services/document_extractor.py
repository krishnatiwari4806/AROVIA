"""Document Extractor and Sanitization Service for PDF and DOCX files."""

import asyncio
import io
import re
import zipfile
from typing import Tuple

import docx
import pdfplumber

from app.core.exceptions import ValidationError

# Magic byte signatures
PDF_MAGIC_BYTES = b"%PDF-"
ZIP_MAGIC_BYTES = b"PK\x03\x04"
MIME_PDF = "application/pdf"
MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MIN_EXTRACTED_CHARACTERS = 50
MAX_EXTRACTED_CHARACTERS = 30000


def verify_file_magic_bytes(content: bytes, filename: str) -> str:
    """Validate file binary signature and verify it matches the declared file extension.

    Args:
        content: Raw bytes of the uploaded file.
        filename: Original filename from client.

    Returns:
        The verified MIME type string.

    Raises:
        ValidationError: If magic bytes do not match or extension is spoofed.
    """
    if not content:
        raise ValidationError("Uploaded file is empty.")

    lower_filename = filename.lower().strip()

    if content.startswith(PDF_MAGIC_BYTES):
        if not lower_filename.endswith(".pdf"):
            raise ValidationError(
                "File signature indicates PDF, but filename extension does not match .pdf"
            )
        return MIME_PDF

    if content.startswith(ZIP_MAGIC_BYTES):
        if not lower_filename.endswith(".docx"):
            raise ValidationError(
                "File signature indicates DOCX/ZIP, but filename extension does not match .docx"
            )
        # Verify it is a valid DOCX zip containing standard Word XML parts
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                file_list = zf.namelist()
                has_content_types = "[Content_Types].xml" in file_list
                has_word_doc = any(name.startswith("word/") for name in file_list)
                if not (has_content_types and has_word_doc):
                    raise ValidationError(
                        "File is a valid ZIP archive but not a valid Word (.docx) document."
                    )
        except zipfile.BadZipFile:
            raise ValidationError("Corrupted DOCX file structure.")
        return MIME_DOCX

    raise ValidationError(
        "Invalid file format or corrupted signature. Only standard text-based PDF and DOCX files are allowed."
    )


def _sync_extract_pdf(content: bytes) -> str:
    """Synchronously extract plain text from PDF bytes using pdfplumber."""
    text_chunks: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n\n".join(text_chunks)


def _sync_extract_docx(content: bytes) -> str:
    """Synchronously extract plain text from DOCX bytes using python-docx."""
    doc = docx.Document(io.BytesIO(content))
    text_chunks: list[str] = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            text_chunks.append(para.text.strip())

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                text_chunks.append(row_text)

    return "\n".join(text_chunks)


def sanitize_extracted_text(text: str) -> str:
    """Sanitize extracted plain text: strip control chars and bound length."""
    # Strip null bytes and non-printable control characters (except newline, tab, carriage return)
    sanitized = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    # Normalize excessive newlines
    sanitized = re.sub(r"\n{3,}", "\n\n", sanitized).strip()
    # Bound to maximum characters
    if len(sanitized) > MAX_EXTRACTED_CHARACTERS:
        sanitized = sanitized[:MAX_EXTRACTED_CHARACTERS]
    return sanitized


async def extract_and_sanitize_text(content: bytes, filename: str) -> Tuple[str, str]:
    """Verify magic bytes, safely extract text in threadpool, sanitize, and validate character count.

    Args:
        content: Raw bytes of uploaded file.
        filename: Original filename.

    Returns:
        Tuple of (sanitized_text, verified_mime_type).

    Raises:
        ValidationError: If format is invalid, extraction fails, or text is under 50 characters.
    """
    mime_type = verify_file_magic_bytes(content, filename)

    try:
        if mime_type == MIME_PDF:
            raw_text = await asyncio.to_thread(_sync_extract_pdf, content)
        elif mime_type == MIME_DOCX:
            raw_text = await asyncio.to_thread(_sync_extract_docx, content)
        else:
            raise ValidationError(f"Unsupported MIME type: {mime_type}")
    except ValidationError:
        raise
    except Exception as exc:
        raise ValidationError(
            f"Failed to extract text from document: {str(exc)}"
        ) from exc

    sanitized = sanitize_extracted_text(raw_text)

    # Check minimum character count for non-whitespace characters
    non_ws_chars = len(re.sub(r"\s+", "", sanitized))
    if non_ws_chars < MIN_EXTRACTED_CHARACTERS:
        raise ValidationError(
            "No extractable text found. Please upload a standard text-based PDF or DOCX resume (scanned image PDFs are not supported)."
        )

    return sanitized, mime_type
